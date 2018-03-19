
#referral.py - installed on teacher computers. Takes an input of the name
#and behavior and spits out several files on the sharedrive record

#ultimate TODO is to alert Evans via email.

import docx, os, openpyxl, datetime



def writereferral(blanks, issues, remedies):

#master excel
    
    wb1 = openpyxl.load_workbook('H:\\Referrals Margate\\MasterReferral.xlsx')
    sheet = wb1.active
    nextrow= sheet.max_row + 1
    sheet.cell(row=nextrow, column=1).value = blanks[0] #name
    sheet.cell(row=nextrow, column=2).value = 'Handwerg'

    dt = datetime.datetime.now()
    
    timedate = ('%s-%s-%s %s-%s' % (dt.month, dt.day, dt.year, dt.hour, dt.second))
    sheet.cell(row=nextrow, column=3).value = timedate
    day = ('%s-%s-%s' % (dt.month, dt.day, dt.year))

    behaviors = ', '.join(issues)
    sheet.cell(row=nextrow, column=4).value = behaviors #behaviors

    sheet.cell(row=nextrow, column=5).value = blanks[1] #behavior comments
    
    corrections = ', '.join(remedies)
    sheet.cell(row=nextrow, column=6).value = corrections #remedies

    sheet.cell(row=nextrow, column=7).value = blanks[2] #other remedy
      
    
    wb1.save('H:\\Referrals Margate\\MasterReferral.xlsx')
    try:
        wb1.save('H:\\Referrals Margate\\MasterReferralAccessCopy.xlsx')#makes a copy for Evans to use
    except Exception as err:
        print(err)
    wb1.close()

#individual excel
    if ('%s.xlsx' % blanks[0]) not in os.listdir('H:\\Referrals Margate\\Individual Files'):
        wb=openpyxl.Workbook() #makes a file on a first timer
        wb.save('H:\\Referrals Margate\\Individual Files\\%s.xlsx' % blanks[0])
        

    wb1 = openpyxl.load_workbook('H:\\Referrals Margate\\Individual Files\\%s.xlsx' % blanks[0])
    sheet = wb1.active
    nextrow= sheet.max_row + 1
    
    sheet.cell(row=nextrow, column=1).value = blanks[0] #name
    sheet.cell(row=nextrow, column=2).value = 'Handwerg'
    sheet.cell(row=nextrow, column=3).value = timedate
    behaviors = ', '.join(issues)
    sheet.cell(row=nextrow, column=4).value = behaviors #behaviors

    sheet.cell(row=nextrow, column=5).value = blanks[1] #behavior comments
    
    corrections = ', '.join(remedies)
    sheet.cell(row=nextrow, column=6).value = corrections #remedies

    sheet.cell(row=nextrow, column=7).value = blanks[2] #other remedy
    wb1.save('H:\\Referrals Margate\\Individual Files\\%s.xlsx' % blanks[0])
    wb1.close()
#word file

    
    doc=docx.Document('H:\Referrals Margate\\Referral Original.docx')
    doc.save('H:\\Referrals Margate\\Word Files\\%s %s.docx' % (blanks[0], timedate))
    doc=docx.Document('H:\\Referrals Margate\\Word Files\\%s %s.docx' % (blanks[0], timedate))
    doc.paragraphs[3].runs[2].text = blanks[0]
    doc.paragraphs[3].runs[2].underline = True
    doc.paragraphs[5].runs[0].text= "Referring Staff:"
    doc.paragraphs[5].runs[1].text= ' Handwerg'
    doc.paragraphs[5].runs[1].underline = True
    doc.paragraphs[5].runs[6].text= 'Date and Time:'
    doc.paragraphs[5].runs[7].text= timedate
    doc.paragraphs[5].runs[7].underline = True
    doc.paragraphs[15].runs[2].text= 'Classroom:'
    doc.paragraphs[15].runs[3].text= ' 109'
    doc.paragraphs[15].runs[3].underline= True
    doc.paragraphs[20].runs[0].text = behaviors
    doc.paragraphs[20].runs[0].underline= True
    doc.paragraphs[20].runs[0].bold= False
    doc.paragraphs[24].runs[0].text= corrections
    doc.paragraphs[24].runs[0].underline= True
    doc.paragraphs[24].runs[3].text= 'Other:%s' % (blanks[2])
##    doc.paragraphs[24].runs[2].text= blanks[2]
    doc.paragraphs[24].runs[3].underline= True
    doc.paragraphs[27].runs[0].text= blanks[1]
    doc.paragraphs[27].runs[0].underline= True
    
    #todo erase time from original form
    

    doc.save('H:\\Referrals Margate\\Word Files\\%s %s.docx' % (blanks[0], timedate))


from tkinter import *

def list_of_entries():
    
    a = []
    b =[]
    c =[]

    b.append(e1.get())
    if var1.get() == 1:
        a.append('Inappropriate Language')
   

    if var2.get() == 1:
        a.append('Defiance/Disrespect')


    if var3.get() == 1:
        a.append('Fighting')


    if var4.get() == 1:
        a.append('Not Following Rules')


    if var5.get() == 1:
        a.append('Cell Phone')


    if var6.get() == 1:
        a.append('Disruption Major')


    if var7.get() == 1:
        a.append('Vandalism Major')
 

    if var8.get() == 1:
        a.append('Disruption Minor')

    if var9.get() == 1:
        a.append('Vandalism Minor')

    if var10.get() == 1:
        a.append('Inappropriate use of Technology')   

    b.append(behaviorbox.get("1.0",'end-1c'))

    if var11.get() == 1:
        c.append('Parent Conference')


    if var12.get() == 1:
        c.append('Schedule Change')


    if var13.get() == 1:
        c.append('Referred to Guidance')


    if var14.get() == 1:
        c.append('Suspension')


    if var15.get() == 1:
        c.append('Sent Home')

  
    b.append(e2.get())
##    print(entrylist)
    writereferral(b, a, c)
    print ('done')
    
    
    


master= Tk()

var1 = IntVar()
var2 = IntVar()
var3 = IntVar()
var4 = IntVar()
var5 = IntVar()
var6 = IntVar()
var7 = IntVar()
var8 = IntVar()
var9 = IntVar()
var10 = IntVar()
var11 = IntVar()
var12 = IntVar()
var13 = IntVar()
var14 = IntVar()
var15 = IntVar()

e1 = Entry(master)
e1.grid(row=1, column=0)

Label(master, text="Name").grid(row=0, sticky = W)


c1 = Checkbutton(master,text="Inappropriate Language", variable=var1).grid(row=2, column = 0, sticky=W)
c2 = Checkbutton(master, text="Defiance/Disrespect", variable=var2).grid(row=2, column = 1, sticky=W)
c3 = Checkbutton(master,text="Fighting", variable=var3).grid(row=3, column = 0, sticky=W)
c4 = Checkbutton(master, text="Not Following Rules", variable=var4).grid(row=3, column = 1, sticky=W)
c5 = Checkbutton(master,text="Cell Phone", variable=var5).grid(row=4, column = 0, sticky=W)
c6 = Checkbutton(master, text="Disruption Major", variable=var6).grid(row=4, column = 1, sticky=W)
c7 = Checkbutton(master,text="Vandalism Major", variable=var7).grid(row=5, column = 0, sticky=W)
c8 = Checkbutton(master, text="Disruption Minor", variable=var8).grid(row=5, column = 1, sticky=W)
c9 = Checkbutton(master, text="Vandalism Minor", variable=var9).grid(row=6, column = 0, sticky=W)
c10 = Checkbutton(master, text="Inappropriate Use of Technology", variable=var10).grid(row=6, column = 1, sticky=W)


Label(master, text="Description").grid(row=7, sticky = W)
behaviorbox = Text(master, height = 10, width =30)
behaviorbox.grid(row=8, column = 0, sticky=W, columnspan=2)


Label(master, text="Remedy").grid(row=9, sticky = W)

c11 = Checkbutton(master,text="Parent Conference", variable=var11).grid(row=10, column = 0, sticky=W)
c12 = Checkbutton(master, text="Schedule Change", variable=var12).grid(row=10, column = 1, sticky=W)
c13 = Checkbutton(master,text="Referred to Guidance", variable=var13).grid(row=11, column = 0, sticky=W)
c14 = Checkbutton(master, text="Suspension", variable=var14).grid(row=11, column = 1, sticky=W)
c15 = Checkbutton(master,text="Sent Home", variable=var15).grid(row=12, column = 0, sticky=W)

Label(master, text="Other").grid(row=16, sticky = W)
e2 = Entry(master)
e2.grid(row=17, column=0)


Button(master, text='Submit', command=list_of_entries).grid(row=18, column=0, sticky=W, pady=4)



mainloop()
